"""
Anonymisation du contenu des pièces jointes uploadées.

Remplace les variantes de noms (depuis MatriculeMap) par leur matricule (KU001, etc.)
dans les formats textuels supportés. Pour les formats non supportés (PDF, images),
le fichier est stocké tel quel mais marqué non anonymisé.
"""
import io
import re


SUPPORTED_TEXT_EXT = {'.txt', '.csv', '.md', '.json', '.log'}
SUPPORTED_EXCEL_EXT = {'.xlsx', '.xlsm'}
SUPPORTED_DOCX_EXT = {'.docx'}


def _generate_inverted_variants(variants_to_matricule: dict) -> dict:
    """
    Pour chaque variante "NOM Prénom" (première partie majuscule), ajoute
    aussi "Prénom NOM" qui est l'ordre courant dans les textes libres.
    """
    extra = {}
    for variant, matricule in list(variants_to_matricule.items()):
        parts = variant.split()
        if len(parts) < 2:
            continue
        nom_parts = [p for p in parts if p.isupper() and len(p) > 1]
        prenom_parts = [p for p in parts if not (p.isupper() and len(p) > 1)]
        if nom_parts and prenom_parts:
            inverted = ' '.join(prenom_parts + nom_parts)
            if inverted != variant and inverted not in variants_to_matricule:
                extra[inverted] = matricule
    return extra


def _build_replacement_table():
    """
    Charge la table de correspondance depuis la BDD.
    Retourne une liste de tuples (variant, matricule) triée par longueur décroissante.
    """
    from .models import MatriculeMap

    variants_to_matricule = {}
    for m in MatriculeMap.objects.all():
        for variant in m.variants:
            v = variant.strip()
            if v and len(v) >= 4:
                variants_to_matricule[v] = m.matricule

    # Ajoute les variantes inversées "Prénom NOM"
    variants_to_matricule.update(_generate_inverted_variants(variants_to_matricule))

    return sorted(
        variants_to_matricule.items(),
        key=lambda kv: -len(kv[0])
    )


def _replace_text(text: str, replacements) -> tuple[str, int]:
    """Applique toutes les substitutions sur un texte. Retourne (nouveau_texte, nb_remplacements)."""
    if not text:
        return text, 0
    total = 0
    for variant, matricule in replacements:
        pattern = re.compile(re.escape(variant), flags=re.IGNORECASE)
        new_text, n = pattern.subn(matricule, text)
        if n:
            text = new_text
            total += n
    return text, total


def anonymize_bytes(content: bytes, filename: str, content_type: str = '') -> tuple[bytes, bool, int]:
    """
    Anonymise le contenu d'un fichier en fonction de son type.

    Args:
        content: bytes du fichier uploadé
        filename: nom de fichier d'origine (pour détection extension)
        content_type: type MIME (facultatif)

    Returns:
        (new_content, was_anonymized, nb_replacements)
        - new_content: bytes anonymisés (ou identiques si format non supporté)
        - was_anonymized: True si le fichier a été traité
        - nb_replacements: nombre de remplacements effectués
    """
    lname = (filename or '').lower()
    ext = lname[lname.rfind('.'):] if '.' in lname else ''

    replacements = _build_replacement_table()
    if not replacements:
        return content, False, 0

    # Fichiers texte
    if ext in SUPPORTED_TEXT_EXT:
        try:
            text = content.decode('utf-8', errors='replace')
        except Exception:
            return content, False, 0
        new_text, nb = _replace_text(text, replacements)
        return new_text.encode('utf-8'), True, nb

    # Fichiers Excel
    if ext in SUPPORTED_EXCEL_EXT:
        try:
            from openpyxl import load_workbook
            wb = load_workbook(io.BytesIO(content), data_only=False)
            total = 0
            for sheet in wb.worksheets:
                for row in sheet.iter_rows():
                    for cell in row:
                        if isinstance(cell.value, str):
                            new_val, n = _replace_text(cell.value, replacements)
                            if n:
                                cell.value = new_val
                                total += n
            out = io.BytesIO()
            wb.save(out)
            return out.getvalue(), True, total
        except ImportError:
            return content, False, 0
        except Exception:
            return content, False, 0

    # Fichiers Word
    if ext in SUPPORTED_DOCX_EXT:
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            total = 0
            for para in doc.paragraphs:
                for run in para.runs:
                    if run.text:
                        new_text, n = _replace_text(run.text, replacements)
                        if n:
                            run.text = new_text
                            total += n
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if run.text:
                                    new_text, n = _replace_text(run.text, replacements)
                                    if n:
                                        run.text = new_text
                                        total += n
            out = io.BytesIO()
            doc.save(out)
            return out.getvalue(), True, total
        except ImportError:
            return content, False, 0
        except Exception:
            return content, False, 0

    # Format non supporté (PDF, images, etc.) : stocké tel quel, non anonymisé
    return content, False, 0
